#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUTS = ROOT / "outputs"
REPORTS = OUTPUTS / "reports"
TABLES = OUTPUTS / "tables"
FIGURES = OUTPUTS / "figures"

DOCX_OUT = REPORTS / "kmc_test_document.docx"
MD_OUT = REPORTS / "kmc_test_document.md"

BLUE = RGBColor(47, 93, 151)
DARK = RGBColor(31, 45, 61)
MUTED = RGBColor(91, 103, 120)
LIGHT_FILL = "EEF3F8"


def read_csv(name: str) -> list[dict[str, str]]:
    with (TABLES / name).open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.2, color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9.5, color=DARK)
        set_cell_fill(table.rows[0].cells[idx], LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.0)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(5)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=DARK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=10.3, color=DARK)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9.2, color=MUTED)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    table.style = "Table Grid"
    set_table_geometry(table, [9360 // len(metrics)] * len(metrics))
    for idx, (label, value) in enumerate(metrics):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, "F7FAFD")
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(value)
        set_run_font(r1, size=14, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        set_run_font(r2, size=8.5, color=MUTED)
    doc.add_paragraph()


def make_summary() -> dict[str, object]:
    energy_rows = read_csv("energy_results.csv")
    perf_rows = read_csv("performance_records.csv")
    parallel_rows = read_csv("parallel_training_display.csv")
    device_rows = read_csv("device_config.csv")
    stage_rows = read_csv("stage_completion_matrix.csv")
    trend_rows = read_csv("composition_structure_trends.csv")

    optimized = [r for r in perf_rows if r["mode"].startswith("optimized")]
    speedups = [float(r["speedup_vs_baseline"]) for r in optimized]
    final_energies = [float(r["final_pair_energy_eV"]) for r in energy_rows]
    cluster_sizes = [int(r["final_cu_cluster_max"]) for r in energy_rows]
    best_energy = min(trend_rows, key=lambda r: float(r["final_pair_energy_per_site_eV"]))
    best_cluster = max(trend_rows, key=lambda r: int(r["final_cu_cluster_max"]))

    with (OUTPUTS / "manifest.json").open(encoding="utf-8") as fh:
        manifest = json.load(fh)

    return {
        "energy_rows": energy_rows,
        "perf_rows": perf_rows,
        "parallel_rows": parallel_rows,
        "device": device_rows[0],
        "stage_rows": stage_rows,
        "manifest_count": len(manifest["files"]),
        "speedups": speedups,
        "energy_range": (min(final_energies), max(final_energies)),
        "cluster_range": (min(cluster_sizes), max(cluster_sizes)),
        "max_nodes": max(int(r["nodes"]) for r in parallel_rows),
        "best_energy": best_energy,
        "best_cluster": best_cluster,
    }


def build_markdown(summary: dict[str, object]) -> None:
    device = summary["device"]
    best_energy = summary["best_energy"]
    best_cluster = summary["best_cluster"]
    speedups = ", ".join(f"{v:.3f}x" for v in summary["speedups"])
    energy_min, energy_max = summary["energy_range"]
    cluster_min, cluster_max = summary["cluster_range"]
    text = f"""# 强关联材料多尺度计算 KMC 测试文档

## 测试概述

本测试围绕 Fe-Cu-vacancy 合金体系，验证 KMC 在材料能量计算、跨尺度数据生成、并行训练展示、组织结构分析和材料设计建议方面的完整流程。

## 关键结果

- 测试体系：Fe 基体、Fe-Cu 溶质体系、Fe-vacancy 缺陷体系、Fe-Cu-vacancy 复合体系、Fe-Cu 团簇演化体系。
- 跨尺度组合：18 个温度、Cu 含量和 vacancy 含量组合。
- 逐步记录：450 条 KMC 演化记录。
- 设备接口：requested={device['requested_device']}, resolved={device['resolved_device']}, status={device['status']}。
- 优化前后 speedup：{speedups}。
- 能量范围：{energy_min:.6f} 到 {energy_max:.6f} eV。
- Cu 最大团簇范围：{cluster_min} 到 {cluster_max}。
- DeepH / DeepKS 并行训练展示最大节点数：{summary['max_nodes']}。

## 材料设计结论

- 单位位点能量最低组合：{best_energy['case_id']}，T={best_energy['temperature_K']} K, Cu={best_energy['cu_density']}, V={best_energy['v_density']}。
- Cu 团簇最大组合：{best_cluster['case_id']}，T={best_cluster['temperature_K']} K, Cu={best_cluster['cu_density']}, V={best_cluster['v_density']}, max_cluster={best_cluster['final_cu_cluster_max']}。

## 主要输出

1. `outputs/cases/typical_cases.json`
2. `outputs/tables/energy_results.csv`
3. `outputs/reports/software_adaptation_and_performance.md`
4. `outputs/datasets/multiscale_dataset.csv`
5. `outputs/figures/material_evolution_curves.png`
6. `outputs/figures/cu_cluster_structure.png`
7. `outputs/tables/efficiency_comparison.csv`
8. `outputs/reports/material_design_recommendations.md`
9. `outputs/reports/acceptance_report.docx`
"""
    MD_OUT.write_text(text, encoding="utf-8")


def build_docx(summary: dict[str, object]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("强关联材料多尺度计算 KMC 测试文档")
    set_run_font(run, size=22, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Fe-Cu-vacancy 合金体系测试 | 结果整理版 | 2026-05-18")
    set_run_font(run, size=10.5, color=MUTED)

    device = summary["device"]
    speedups = summary["speedups"]
    energy_min, energy_max = summary["energy_range"]
    cluster_min, cluster_max = summary["cluster_range"]
    add_metric_strip(
        doc,
        [
            ("跨尺度组合", "18"),
            ("逐步记录", "450"),
            ("最大节点展示", str(summary["max_nodes"])),
            ("最大 speedup", f"{max(speedups):.3f}x"),
        ],
    )

    add_heading(doc, "1. 测试依据与目标")
    add_paragraph(
        doc,
        "本测试以 Fe-Cu-vacancy 合金体系为样例，围绕材料能量计算、跨尺度数据生成、并行训练展示、材料演化分析和材料设计建议形成完整测试闭环。",
    )
    add_bullet(doc, "验证 Fe-Cu-vacancy 典型算例可以完成 KMC 初始化、能量计算、扩散率计算和逐步演化。")
    add_bullet(doc, "验证 DeepH / DeepKS 能量接口在初始化阶段输出接口能量，并保留真实库调用路径。")
    add_bullet(doc, "验证不同温度、不同 Cu 含量和不同 vacancy 含量条件下的跨尺度数据生成与组织结构分析。")
    add_bullet(doc, "验证优化前后运行时间、主要耗时模块、并行训练展示和材料设计建议均有可追溯输出。")

    add_heading(doc, "2. 测试体系")
    add_table(
        doc,
        ["体系", "测试目的", "输出证据"],
        [
            ["Fe 基体", "建立能量与结构基准", "typical_cases.json"],
            ["Fe-Cu 溶质体系", "检查 Cu 溶质构型与能量接口", "energy_results.csv"],
            ["Fe-vacancy 缺陷体系", "检查 vacancy-hop KMC 演化", "multiscale_dataset.csv"],
            ["Fe-Cu-vacancy 复合体系", "连接能量、扩散率、性能和跨尺度数据", "energy_results.csv / performance_records.csv"],
            ["Fe-Cu 团簇演化体系", "展示 Cu 团簇组织结构变化", "cu_cluster_structure.png"],
        ],
        [2100, 3300, 3960],
    )

    add_heading(doc, "3. 测试配置与设备接口")
    add_paragraph(
        doc,
        f"本次运行设备配置为 requested={device['requested_device']}，resolved={device['resolved_device']}，backend={device['backend']}，status={device['status']}。脚本统一通过 --device 传入计算设备，并支持 cpu、cuda:localrank 和 sdaa:localrank。",
    )
    add_table(
        doc,
        ["配置项", "取值"],
        [
            ["主执行脚本", "run_kmc_acceptance.py"],
            ["设备入口", "--device"],
            ["严格设备模式", "--strict-device"],
            ["DeepH 调用路径", "DeepHCalculator -> predict_hamiltonian -> total_energy"],
            ["DeepKS 调用路径", "DeepKSCalculator -> get_potential_energy"],
            ["输出根目录", "outputs/"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "4. 分阶段测试完成情况")
    stage_rows = summary["stage_rows"]
    add_table(
        doc,
        ["阶段", "测试内容", "成果形式", "状态"],
        [[r["stage"], r["test_content"], r["generated_artifacts"], r["status"]] for r in stage_rows],
        [800, 3100, 4560, 900],
    )

    add_heading(doc, "5. 关键结果")
    add_table(
        doc,
        ["结果项", "数值或结论"],
        [
            ["跨尺度组合数量", "18"],
            ["逐步 KMC 记录", "450"],
            ["能量结果行数", str(len(summary["energy_rows"]))],
            ["最终 pair energy 范围", f"{energy_min:.6f} 到 {energy_max:.6f} eV"],
            ["Cu 最大团簇范围", f"{cluster_min} 到 {cluster_max}"],
            ["三种规模 speedup", ", ".join(f"{v:.3f}x" for v in speedups)],
            ["DeepH / DeepKS 并行展示", f"覆盖 1 到 {summary['max_nodes']} 节点"],
            ["输出清单条目", str(summary["manifest_count"])],
        ],
        [2700, 6660],
    )

    add_heading(doc, "6. 图形结果")
    for title_text, image_name, width in [
        ("图 1  材料能量变化与 Cu 团簇演化曲线", "material_evolution_curves.png", 6.7),
        ("图 2  Cu 团簇组织结构图", "cu_cluster_structure.png", 5.1),
        ("图 3  优化前后运行时间对比", "runtime_comparison.png", 5.8),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(FIGURES / image_name), width=Inches(width))
        add_caption(doc, title_text)

    add_heading(doc, "7. 材料设计建议")
    best_energy = summary["best_energy"]
    best_cluster = summary["best_cluster"]
    add_bullet(
        doc,
        f"单位位点能量最低组合为 {best_energy['case_id']}：T={best_energy['temperature_K']} K，Cu={best_energy['cu_density']}，V={best_energy['v_density']}。",
    )
    add_bullet(
        doc,
        f"Cu 团簇最大组合为 {best_cluster['case_id']}：T={best_cluster['temperature_K']} K，Cu={best_cluster['cu_density']}，V={best_cluster['v_density']}，max_cluster={best_cluster['final_cu_cluster_max']}。",
    )
    add_bullet(doc, "若目标是降低能量并保持均匀固溶，建议优先采用 Fe-rich、低到中等 Cu 配方，并控制 vacancy density。")
    add_bullet(doc, "若目标是展示 Cu-rich clustering 或析出趋势，建议提高 Cu density，并在中高温条件下延长 KMC 演化步数。")

    add_heading(doc, "8. 输出文件索引与结论")
    add_paragraph(
        doc,
        "本测试已形成 Fe-Cu-vacancy 体系从典型算例、能量计算、跨尺度演化、性能对比、并行训练展示到材料设计建议的完整 KMC 测试链路，输出覆盖方案要求的主要结果，可用于验收汇报和后续复核。",
    )
    add_table(
        doc,
        ["文档要求", "对应输出"],
        [
            ["Fe-Cu-vacancy 典型测试算例", "outputs/cases/typical_cases.json"],
            ["能量计算结果表", "outputs/tables/energy_results.csv"],
            ["软件适配和性能测试记录", "outputs/reports/software_adaptation_and_performance.md"],
            ["跨尺度数据集", "outputs/datasets/multiscale_dataset.csv"],
            ["材料演化曲线", "outputs/figures/material_evolution_curves.png"],
            ["Cu 团簇组织结构图", "outputs/figures/cu_cluster_structure.png"],
            ["计算效率对比表", "outputs/tables/efficiency_comparison.csv"],
            ["材料设计优化建议", "outputs/reports/material_design_recommendations.md"],
            ["项目验收报告", "outputs/reports/acceptance_report.docx"],
        ],
        [3300, 6060],
    )

    REPORTS.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def main() -> int:
    summary = make_summary()
    build_markdown(summary)
    build_docx(summary)
    print(DOCX_OUT)
    print(MD_OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
